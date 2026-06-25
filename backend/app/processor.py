from datetime import datetime
from io import BytesIO, StringIO
from numbers import Number
from pathlib import Path
from typing import Any, Dict, List, Optional
import json
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import plotly.express as px
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi import UploadFile


UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


class DataProcessor:
    
    def __init__(self):
        self.data: Optional[pd.DataFrame] = None
        self.filename: Optional[str] = None

    async def load_file(self, file: UploadFile) -> pd.DataFrame:
        if not file.filename:
            raise ValueError("Не указано имя файла")

        content = await file.read()

        if not content:
            raise ValueError("Загруженный файл пуст")

        extension = Path(file.filename).suffix.lower()

        if extension in {".csv", ".txt"}:
            data = self._read_text_file(content)

        elif extension in {".xlsx", ".xls"}:
            data = pd.read_excel(BytesIO(content))

        else:
            raise ValueError(
                f"Формат {extension} не поддерживается"
            )

        if data.empty:
            raise ValueError("В файле нет данных")

        
        data.columns = [
            str(column).strip()
            for column in data.columns
        ]

        self.data = data
        self.filename = file.filename

        return data
# цсвэшки не всегда в утф 8 сохраняются
    @staticmethod
    def _read_text_file(content: bytes) -> pd.DataFrame:
        encodings = ("utf-8-sig", "utf-8", "cp1251")# добавила запасной вариант

        for encoding in encodings:
            try:
                text = content.decode(encoding)

                return pd.read_csv(
                    StringIO(text),
                    sep=None,
                    engine="python"
                )

            except UnicodeDecodeError:
                continue

            except pd.errors.ParserError as error:
                raise ValueError(
                    f"Не удалось разобрать структуру файла: {error}"
                )

        raise ValueError(
            "Не удалось определить кодировку файла"
        )

    def _get_data(self) -> pd.DataFrame:
        if self.data is None:
            raise ValueError("Сначала необходимо загрузить файл")

        return self.data

    def get_preview(
        self,
        n_rows: int = 10
    ) -> List[Dict[str, Any]]:
      
        if self.data is None:
            return []

        # to_json преобразует NaN в null и корректно обрабатывает даты
        preview_json = self.data.head(n_rows).to_json(
            orient="records",
            date_format="iso"
        )

        return json.loads(preview_json)

    def get_columns_info(self) -> Dict[str, List[str]]:
        if self.data is None:
            return {
                "columns": [],
                "numeric_columns": []
            }

        numeric_columns = (
            self.data
            .select_dtypes(include=[np.number])
            .columns
            .tolist()
        )

        return {
            "columns": self.data.columns.tolist(),
            "numeric_columns": numeric_columns
        }

    def calculate_statistics(
        self
    ) -> Dict[str, Dict[str, Optional[float]]]:
        data = self._get_data()

        numeric_columns = data.select_dtypes(
            include=[np.number]
        ).columns

        statistics = {}

        for column in numeric_columns:
            values = data[column].dropna()

            if values.empty:
                continue

            std = values.std()

            statistics[column] = {
                "mean": float(values.mean()),
                "std": None if pd.isna(std) else float(std),
                "min": float(values.min()),
                "max": float(values.max()),
                "median": float(values.median()),
                "q1": float(values.quantile(0.25)),
                "q3": float(values.quantile(0.75))
            }

        return statistics

    def create_plot(
        self,
        x_col: str,
        y_col: str,
        plot_type: str = "line",
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        data = self._get_data()

        if y_col not in data.columns:
            raise ValueError(
                f"Столбец «{y_col}» не найден"
            )

        if plot_type != "box" and x_col not in data.columns:
            raise ValueError(
                f"Столбец «{x_col}» не найден"
            )

        if plot_type == "line":
            default_title = f"Зависимость {y_col} от {x_col}"

            figure = px.line(
                data,
                x=x_col,
                y=y_col,
                title=title or default_title
            )

        elif plot_type == "scatter":
            default_title = (
                f"Диаграмма рассеяния {y_col} от {x_col}"
            )

            figure = px.scatter(
                data,
                x=x_col,
                y=y_col,
                title=title or default_title
            )

        elif plot_type == "bar":
            default_title = (
                f"Столбчатая диаграмма {y_col} от {x_col}"
            )

            figure = px.bar(
                data,
                x=x_col,
                y=y_col,
                title=title or default_title
            )

        elif plot_type == "box":
            default_title = (
                f"Распределение значений {y_col}"
            )

            figure = px.box(
                data,
                y=y_col,
                title=title or default_title
            )

        else:
            raise ValueError(
                f"Неизвестный тип графика: {plot_type}"
            )

        figure.update_layout(
            template="plotly_white",
            xaxis_title=x_col if plot_type != "box" else None,
            yaxis_title=y_col,
            hovermode="closest"
        )

        return json.loads(figure.to_json())

    def generate_report_word(
        self,
        report_data: Dict[str, Any]
    ) -> str:
        data = self._get_data()
        document = Document()

        title = document.add_heading(
            f"Лабораторная работа: {report_data['title']}",
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_student_information(document, report_data)

        document.add_heading("Цель работы", level=1)
        document.add_paragraph(report_data["objective"])

        document.add_heading("Теоретическая часть", level=1)
        document.add_paragraph(report_data["theory"])

        self._add_data_table(document, data)
        self._add_report_plot(document, data)

        document.add_heading("Анализ результатов", level=1)
        document.add_paragraph(report_data["analysis"])

        document.add_heading("Выводы", level=1)
        document.add_paragraph(report_data["conclusions"])

        additional_notes = report_data.get("additional_notes")

        if additional_notes:
            document.add_heading(
                "Дополнительные замечания",
                level=1
            )
            document.add_paragraph(additional_notes)

        filename = (
            f"report_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}.docx"
        )

        file_path = UPLOAD_DIR / filename
        document.save(file_path)

        return str(file_path)

    @staticmethod
    def _add_student_information(
        document: Document,
        report_data: Dict[str, Any]
    ):
        document.add_heading(
            "Информация о студенте",
            level=1
        )

        document.add_paragraph(
            f"ФИО: {report_data['student_name']}"
        )
        document.add_paragraph(
            f"Группа: {report_data['group']}"
        )
        document.add_paragraph(
            f"Дата: {datetime.now().strftime('%d.%m.%Y')}"
        )

    def _add_data_table(
        self,
        document: Document,
        data: pd.DataFrame
    ):
        document.add_heading(
            "Экспериментальные данные",
            level=1
        )

        table = document.add_table(
            rows=1,
            cols=len(data.columns)
        )
        table.style = "Table Grid"

        for index, column in enumerate(data.columns):
            table.rows[0].cells[index].text = str(column)

        for _, row in data.head(20).iterrows():
            cells = table.add_row().cells

            for index, value in enumerate(row):
                cells[index].text = self._format_table_value(
                    value
                )

        note = document.add_paragraph()
        note_run = note.add_run(
            f"Показаны первые 20 строк из {len(data)}."
        )
        note_run.italic = True

    @staticmethod
    def _format_table_value(value: Any) -> str:
        if pd.isna(value):
            return ""

        if isinstance(value, Number) and not isinstance(value, bool):
            return f"{float(value):.3f}"

        return str(value)

    @staticmethod
    def _add_report_plot(
        document: Document,
        data: pd.DataFrame
    ):
        document.add_heading("Графики", level=1)

        numeric_columns = data.select_dtypes(
            include=[np.number]
        ).columns.tolist()

        if len(numeric_columns) < 2:
            document.add_paragraph(
                "Для построения графика необходимо "
                "не менее двух числовых столбцов."
            )
            return

        x_column = numeric_columns[0]
        y_column = numeric_columns[1]

        plot_data = data[
            [x_column, y_column]
        ].dropna()

        if plot_data.empty:
            document.add_paragraph(
                "Недостаточно данных для построения графика."
            )
            return

        figure, axes = plt.subplots(figsize=(8, 5))

        axes.plot(
            plot_data[x_column],
            plot_data[y_column],
            marker="o",
            linewidth=1.5
        )

        axes.set_xlabel(x_column)
        axes.set_ylabel(y_column)
        axes.set_title(
            f"Зависимость {y_column} от {x_column}"
        )
        axes.grid(True, alpha=0.3)

        image_buffer = BytesIO()

        figure.savefig(
            image_buffer,
            format="png",
            dpi=200,
            bbox_inches="tight"
        )

        plt.close(figure)
        image_buffer.seek(0)

        picture = document.add_picture(
            image_buffer,
            width=Inches(6)
        )

        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        image_buffer.close()


processor = DataProcessor()

